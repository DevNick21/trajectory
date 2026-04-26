"""Onboarding CV parser (PROCESS Entry 49).

Takes an uploaded CV (PDF, DOCX, or pasted text) and extracts a
structured `CVImport` so the onboarding wizard can pre-fill name,
location, contact email, role rows, education, projects, and skills.

The user reviews/edits the extracted result instead of re-typing what's
already on their CV. The raw text also seeds the WritingStyleProfile
samples downstream — a full CV is ~400+ words of the user's own
writing, far richer than the wizard's 3 free-form paragraphs.

Routing: Sonnet via `call_structured`. Cheap (~$0.05) and fast (~5s).
Untrusted input is shielded with Tier 1 regex first — uploaded CVs
sometimes carry recruiter-injected boilerplate.
"""

from __future__ import annotations

import logging
from io import BytesIO
from typing import Optional

from ..config import settings
from ..llm import call_structured
from ..schemas import CVImport, CVImportLLMOutput
from ..validators.content_shield import tier1 as _shield_tier1

logger = logging.getLogger(__name__)


_SYSTEM_PROMPT = """You are a CV parser. The user will paste the raw
text of their CV (already extracted from PDF or DOCX). Your job is to
return a structured representation that an onboarding wizard can
pre-fill.

Rules:
1. Be faithful — extract what's THERE, don't infer or embellish.
2. Roles are listed in reverse-chronological order on the CV; preserve
   that order in the `roles` array.
3. Each role's `bullets` should be the verbatim achievement bullets
   from that section. Don't paraphrase. Strip leading bullet
   characters (•, -, *, etc.) but keep the wording intact.
4. `name` is the candidate's name from the top of the CV. Often it's
   the largest text. If multiple candidate names appear (e.g. a
   reference's name lower down), the top-of-document one wins.
5. `base_location` is the city the candidate is based in — typically
   in the CV header next to the email/phone. UK cities only ("London",
   "Manchester", etc.) — strip "United Kingdom" or "UK" suffixes.
6. `contact_email` — exactly as written.
7. `professional_summary` — the headline/summary paragraph if present
   (often called "Profile", "About", "Summary"). 1-3 sentences. If
   absent, leave null.
8. `skills` — the bullet/comma list from the Skills section. One
   skill per array entry. Don't merge ("Python, Django" → two entries).
9. `extraction_confidence` (1-10) — your honest assessment of how
   well-structured the input was and how cleanly you parsed it. A
   neat reverse-chronological CV gets 9-10; a messy one-page summary
   gets 4-6; an obviously truncated/garbled extraction gets 1-3.

The `raw_text` field is filled in by the caller, not you. Don't try
to populate it.
"""


async def parse(
    *,
    cv_text: str,
    session_id: Optional[str] = None,
) -> CVImport:
    """Sonnet pass: free-form CV text -> structured CVImport.

    `cv_text` is the post-extraction raw text — caller is responsible
    for PDF/DOCX -> str. The Tier 1 content shield runs here before
    the agent sees any of it (CVs are user input and may carry
    injection attempts pasted from recruiter emails).
    """
    if not cv_text or not cv_text.strip():
        raise ValueError("cv_text is empty")

    shielded = _shield_tier1(cv_text).cleaned_text

    extracted: CVImportLLMOutput = await call_structured(
        agent_name="cv_parser",
        system_prompt=_SYSTEM_PROMPT,
        user_input=shielded,
        output_schema=CVImportLLMOutput,
        model=settings.sonnet_model_id,
        effort="medium",
        session_id=session_id,
    )
    # raw_text is supplied by the caller, not Sonnet — feed the
    # un-shielded original through so style_extractor downstream sees
    # the user's actual writing voice.
    out = CVImport(**extracted.model_dump(), raw_text=cv_text)
    logger.info(
        "cv_parser: name=%r roles=%d education=%d projects=%d "
        "skills=%d confidence=%d",
        out.name, len(out.roles), len(out.education),
        len(out.projects), len(out.skills), out.extraction_confidence,
    )
    return out


# ---------------------------------------------------------------------------
# File-format helpers — caller hands us bytes, we return text
# ---------------------------------------------------------------------------


def _extract_pdf_text(data: bytes) -> str:
    """Extract text from a PDF byte stream. Uses pypdf (already a
    dependency via the renderers chain). For LaTeX-typeset CVs the
    extraction can be jagged — that's fine; Sonnet tolerates noise."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "pypdf is required for PDF CV extraction"
        ) from exc

    reader = PdfReader(BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            logger.warning("pypdf page extract failed: %s", exc)
    return "\n\n".join(p for p in pages if p.strip())


def _extract_docx_text(data: bytes) -> str:
    """Extract text from a DOCX byte stream via python-docx (already
    a renderer dependency). Tables and headers are flattened to
    paragraph order — close enough for Sonnet."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX CV extraction"
        ) from exc

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(*, data: bytes, filename: str) -> str:
    """Dispatch on file extension. Plain-text upload bypasses both."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf_text(data)
    if name.endswith(".docx"):
        return _extract_docx_text(data)
    # .txt / .md / unknown — assume utf-8 text
    return data.decode("utf-8", errors="replace")
