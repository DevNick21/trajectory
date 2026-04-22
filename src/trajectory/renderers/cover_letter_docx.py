"""Cover letter → .docx via python-docx.

UK convention: date top-right, address block, salutation, body paragraphs,
signoff. Citation markers stripped.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..schemas import CoverLetterOutput


def _safe_filename(addressed_to: str) -> str:
    def clean(s: str) -> str:
        return "".join(c if c.isalnum() or c in "-_" else "_" for c in s)

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    return f"CoverLetter_{clean(addressed_to[:40])}_{ts}.docx"


def render_cover_letter_docx(
    cl: CoverLetterOutput,
    out_dir: Path,
    sender_name: str = "",
) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / _safe_filename(cl.addressed_to)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # Date (right-aligned)
    date_p = doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_after = Pt(18)
    for run in date_p.runs:
        run.font.size = Pt(11)

    # Addressee block
    addr_p = doc.add_paragraph(cl.addressed_to)
    addr_p.paragraph_format.space_after = Pt(18)
    for run in addr_p.runs:
        run.font.size = Pt(11)

    # Body paragraphs
    for para in cl.paragraphs:
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Pt(0)
        for run in p.runs:
            run.font.size = Pt(11)

    # Signoff
    doc.add_paragraph()
    signoff_p = doc.add_paragraph("Yours sincerely,")
    signoff_p.paragraph_format.space_after = Pt(36)
    for run in signoff_p.runs:
        run.font.size = Pt(11)

    name_p = doc.add_paragraph(sender_name or "")
    for run in name_p.runs:
        run.bold = True
        run.font.size = Pt(11)

    doc.save(str(out_path))
    return out_path
