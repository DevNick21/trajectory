"""CV → .docx via python-docx.

UK convention: name header, contact row, professional summary,
reverse-chrono experience, education, skills. Citation markers on
CVBullet.citations are stripped — they're internal references, not
part of the delivered document.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..config import settings
from ..schemas import CVOutput


_HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)
_RULE_COLOR = RGBColor(0xCC, 0xCC, 0xCC)


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = _HEADING_COLOR
    _add_horizontal_rule(doc)


def _contact_string(contact: dict) -> str:
    parts = []
    for key in ("email", "phone", "location", "linkedin", "github"):
        if val := contact.get(key):
            parts.append(str(val))
    return "  |  ".join(parts)


def _safe_filename(name: str, company: str = "") -> str:
    def clean(s: str) -> str:
        return "".join(c if c.isalnum() or c in "-_" else "_" for c in s)

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    if company:
        return f"{clean(name)}_CV_{clean(company)}_{ts}.docx"
    return f"{clean(name)}_CV_{ts}.docx"


def render_cv_docx(cv: CVOutput, out_dir: Path, company: str = "") -> Path:
    """Produce a .docx file and return its path."""
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / _safe_filename(cv.name, company)

    doc = Document()

    # Narrow margins for a denser 2-page layout.
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv.name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = _HEADING_COLOR

    # Contact
    contact_p = doc.add_paragraph(_contact_string(cv.contact))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(6)
    for run in contact_p.runs:
        run.font.size = Pt(9)

    # Professional summary
    _section_heading(doc, "Professional Summary")
    summary_p = doc.add_paragraph(cv.professional_summary)
    summary_p.paragraph_format.space_after = Pt(4)
    for run in summary_p.runs:
        run.font.size = Pt(10)

    # Experience
    _section_heading(doc, "Experience")
    for role in cv.experience:
        role_p = doc.add_paragraph()
        role_p.paragraph_format.space_before = Pt(6)
        role_p.paragraph_format.space_after = Pt(0)
        title_run = role_p.add_run(role.title)
        title_run.bold = True
        title_run.font.size = Pt(10)
        role_p.add_run(f"  —  {role.company}").font.size = Pt(10)
        dates_run = role_p.add_run(f"\t{role.dates}")
        dates_run.font.size = Pt(9)
        dates_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        for bullet in role.bullets:
            b_p = doc.add_paragraph(style="List Bullet")
            b_p.paragraph_format.space_after = Pt(1)
            b_run = b_p.add_run(bullet.text)
            b_run.font.size = Pt(10)

    # Education
    _section_heading(doc, "Education")
    for edu in cv.education:
        edu_p = doc.add_paragraph()
        edu_p.paragraph_format.space_before = Pt(4)
        deg = edu.get("degree") or edu.get("qualification") or ""
        inst = edu.get("institution") or edu.get("school") or ""
        years = edu.get("dates") or edu.get("years") or ""
        bold_run = edu_p.add_run(f"{deg}  —  {inst}")
        bold_run.bold = True
        bold_run.font.size = Pt(10)
        if years:
            yr_run = edu_p.add_run(f"\t{years}")
            yr_run.font.size = Pt(9)
            yr_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Skills
    _section_heading(doc, "Skills")
    skills_p = doc.add_paragraph("  ·  ".join(cv.skills))
    skills_p.paragraph_format.space_after = Pt(4)
    for run in skills_p.runs:
        run.font.size = Pt(10)

    # Projects (optional)
    if cv.projects:
        _section_heading(doc, "Projects")
        for proj in cv.projects:
            proj_p = doc.add_paragraph()
            proj_p.paragraph_format.space_before = Pt(4)
            name_r = proj_p.add_run(proj.get("name") or "Project")
            name_r.bold = True
            name_r.font.size = Pt(10)
            desc = proj.get("description") or proj.get("summary") or ""
            if desc:
                proj_p.add_run(f" — {desc}").font.size = Pt(10)

    doc.save(str(out_path))
    return out_path
