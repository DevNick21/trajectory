"""Onboarding CV parser — two tiers.

  - **Tier 0 (`tier0_extract`)** — pure regex / heuristic. Pulls name,
    email, phone, LinkedIn, GitHub, location, and the broad shape of
    role lines from the raw text. No LLM, < 100 ms, deterministic.
    The onboarding endpoint returns this immediately so the wizard
    doesn't sit on "Reading your CV..." for a minute.

  - **Tier 1 (`parse`)** — Haiku via `call_structured`. Cleans up what
    tier-0 couldn't (multi-line bullets, education, projects, skill
    list disambiguation). Runs as a separate optional endpoint
    (`/api/onboarding/cv_enrich`) so the wizard can fire it in the
    background while the user keeps editing the prefilled fields.

The user uploads -> tier-0 returns in ~1s -> they advance through the
wizard -> tier-1 (Haiku, ~5s) finishes in the background and merges
its richer extraction into the draft when ready.

Routing: Tier 1 uses Haiku, not Sonnet (downgraded 2026-05-22).
Structured CV extraction is mechanical reshape, not judgement —
Haiku handles it for ~3x less cost and ~3x less latency.

Untrusted input is shielded with Tier 1 regex before either tier
sees it — uploaded CVs sometimes carry recruiter-injected prompts.
"""

from __future__ import annotations

import logging
import re
from io import BytesIO
from typing import Optional

from ..config import settings
from ..llm import call_structured
from ..schemas import CVImport, CVImportLLMOutput, CVImportRole
from ..validators.content_shield import tier1 as _shield_tier1

logger = logging.getLogger(__name__)


_SYSTEM_PROMPT = """You are a CV parser. The user will paste the raw
text of their CV (already extracted from PDF or DOCX). Your job is to
return a structured representation that an onboarding wizard can
pre-fill, PLUS a 2-3 paragraph chronological career narrative.

Rules for the structured extraction:
1. Be faithful — extract what's THERE, don't infer or embellish.
2. Roles are listed in reverse-chronological order on the CV; preserve
   that order in the `roles` array.
3. Each role's `bullets` should be the verbatim achievement bullets
   from that section. Don't paraphrase. Strip leading bullet
   characters (•, -, *, etc.) but keep the wording intact.
4. `name` is the candidate's name from the top of the CV. Often it's
   the largest text. If multiple candidate names appear (e.g. a
   reference's name lower down), the top-of-document one wins.
5. `base_location` is the city the candidate is based in — typically
   in the CV header next to the email/phone. UK cities only ("London",
   "Manchester", etc.) — strip "United Kingdom" or "UK" suffixes.
6. `contact_email` — exactly as written.
7. `professional_summary` — the headline/summary paragraph if present
   (often called "Profile", "About", "Summary"). 1-3 sentences. If
   absent, leave null.
8. `skills` — the bullet/comma list from the Skills section. One
   skill per array entry. Don't merge ("Python, Django" → two entries).
9. `extraction_confidence` (1-10) — your honest assessment of how
   well-structured the input was and how cleanly you parsed it. A
   neat reverse-chronological CV gets 9-10; a messy one-page summary
   gets 4-6; an obviously truncated/garbled extraction gets 1-3.

Rules for `narrative`:
10. 2-3 paragraphs, 100-200 words total.
11. CHRONOLOGICAL — start with earliest education/role, end with the
    most recent. The CV is in reverse-chron; you reverse it back.
12. Third person is fine; first person is fine too if their
    professional_summary uses it.
13. Each role gets one sentence describing scope + a defining bullet.
    Don't enumerate every bullet — pick the most concrete one.
14. Honest. A 2-month contract is a 2-month contract; don't inflate.
15. No clichés. Banned: "passionate", "results-driven", "proven track
    record", "leverage" (as a verb). The self-audit downstream will
    flag these so don't ship them.
16. When the CV is too sparse to summarise (e.g. one role, no detail),
    set `narrative` to null. Don't fabricate.

The `raw_text` field is filled in by the caller, not you. Don't try
to populate it.
"""


# ---------------------------------------------------------------------------
# Tier 0 — regex + heuristic. No LLM. ~50ms.
# ---------------------------------------------------------------------------


_EMAIL_RE = re.compile(r"\b[\w.+-]+@[\w-]+(?:\.[\w-]+)+\b")
# UK + international phone numbers. Anchored on +44 or a 9-15 digit
# run with optional spaces/dashes/parens; lower bound 9 digits to avoid
# 8-digit CRNs and other false matches.
_PHONE_RE = re.compile(
    r"\+?\(?\d[\d\s\-\(\)]{8,18}\d"
)
_LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.IGNORECASE)
_GITHUB_RE = re.compile(r"github\.com/[\w\-]+", re.IGNORECASE)
# A CV's name is usually 2-4 capitalised words on the first non-empty
# line. Allow accents, hyphens, apostrophes, periods (e.g. "K. Iheanacho").
_NAME_LINE_RE = re.compile(
    r"^([A-Z][A-Za-z\.'-]+(?:\s+[A-Z][A-Za-z\.'-]+){1,4})\s*$"
)
# Recognised UK cities + a couple of EU near-misses. The list is short
# on purpose — we just need a strong signal that THIS LINE has a city
# in it. Misses fall through to the LLM tier without losing data.
_UK_CITIES = {
    "london", "manchester", "birmingham", "leeds", "liverpool",
    "bristol", "edinburgh", "glasgow", "sheffield", "newcastle",
    "cardiff", "belfast", "nottingham", "southampton", "leicester",
    "coventry", "bradford", "stoke", "wolverhampton", "plymouth",
    "derby", "swansea", "milton keynes", "oxford", "cambridge",
    "brighton", "york", "reading", "preston", "aberdeen",
}
# "Senior X at Acme" / "Senior X | Acme" / "Senior X — Acme (2024)"
_ROLE_LINE_RE = re.compile(
    r"^\s*([A-Z][A-Za-z0-9 /+&,.()-]{3,60}?)"  # title
    r"\s*(?:at|@|\||—|–|-)\s+"                  # delimiter
    r"([A-Z][A-Za-z0-9 &.,'/-]{1,60}?)"         # company
    r"\s*(?:\(([^)]+)\))?\s*$",                  # optional dates in parens
    re.MULTILINE,
)
# Generic stop tokens for the "name line" heuristic — we don't want a
# header like "CURRICULUM VITAE" or "PROFESSIONAL SUMMARY" treated as
# the candidate's name.
_NAME_STOP_TOKENS = {
    "curriculum", "vitae", "resume", "cv", "professional",
    "summary", "profile", "about", "contact", "experience",
    "education", "skills", "references",
}


def _first_non_blank_lines(text: str, n: int = 8) -> list[str]:
    out: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        out.append(line)
        if len(out) >= n:
            break
    return out


def _looks_like_name(line: str) -> bool:
    """Heuristic — accept a 2-4 word capitalised line without digits."""
    if "@" in line or any(ch.isdigit() for ch in line):
        return False
    if not _NAME_LINE_RE.match(line):
        return False
    tokens = {t.strip(".'").lower() for t in line.split()}
    if tokens & _NAME_STOP_TOKENS:
        return False
    return True


def _extract_name(text: str) -> Optional[str]:
    for line in _first_non_blank_lines(text, n=6):
        if _looks_like_name(line):
            return line
    return None


def _extract_email(text: str) -> Optional[str]:
    match = _EMAIL_RE.search(text)
    return match.group(0) if match else None


def _extract_phone(text: str) -> Optional[str]:
    # Grab the first match in the top of the CV (header section).
    head = "\n".join(_first_non_blank_lines(text, n=20))
    match = _PHONE_RE.search(head)
    return match.group(0).strip() if match else None


def _extract_linkedin(text: str) -> Optional[str]:
    match = _LINKEDIN_RE.search(text)
    return match.group(0) if match else None


def _extract_github(text: str) -> Optional[str]:
    match = _GITHUB_RE.search(text)
    return match.group(0) if match else None


def _extract_location(text: str) -> Optional[str]:
    """Scan the top of the CV for a recognised UK city. Strips trailing
    ", United Kingdom" / ", UK"."""
    for line in _first_non_blank_lines(text, n=10):
        low = line.lower()
        for city in _UK_CITIES:
            if city in low:
                # Title-case the city; strip noise after a separator.
                pretty = city.title()
                if pretty == "Stoke":
                    pretty = "Stoke-on-Trent"
                return pretty
    return None


def _extract_roles_skeleton(text: str) -> list[CVImportRole]:
    """Find role lines via regex. Bullets are left to the LLM tier."""
    roles: list[CVImportRole] = []
    for match in _ROLE_LINE_RE.finditer(text):
        title = match.group(1).strip().rstrip(",")
        company = match.group(2).strip().rstrip(",")
        dates_raw = (match.group(3) or "").strip()
        if len(title) < 4 or len(company) < 2:
            continue
        roles.append(CVImportRole(
            title=title,
            company=company,
            dates=dates_raw or "",
            bullets=[],
        ))
        if len(roles) >= 15:
            break
    return roles


def _extract_skills_inline(text: str) -> list[str]:
    """Find a Skills section and split its first paragraph by commas."""
    # Look for "skills" header (any case) followed by a paragraph.
    # Stop at the next double-newline or all-caps header.
    pattern = re.compile(
        r"(?im)^\s*(?:technical\s+)?skills\s*:?\s*\n+([\s\S]{0,800}?)(?:\n\s*\n|\n[A-Z][A-Z\s]{4,})"
    )
    match = pattern.search(text)
    if not match:
        return []
    chunk = match.group(1)
    candidates: list[str] = []
    for token in re.split(r"[,;\n•]+", chunk):
        token = token.strip().strip("·-*").strip()
        if 1 < len(token) < 60 and not token.startswith("("):
            candidates.append(token)
    # Dedup while preserving order.
    seen: set[str] = set()
    unique: list[str] = []
    for c in candidates:
        key = c.lower()
        if key not in seen:
            seen.add(key)
            unique.append(c)
    return unique[:40]


def tier0_extract(cv_text: str) -> CVImport:
    """Pure-Python CV extraction. No LLM. ~50ms.

    Populates the fields a regex pass can confidently reach. The
    `extraction_confidence` is set to 2 deliberately — the wizard
    surfaces this to nudge the user toward running tier-1 enrichment
    if they want the bullets / education / projects filled too.
    """
    if not cv_text or not cv_text.strip():
        raise ValueError("cv_text is empty")
    return CVImport(
        name=_extract_name(cv_text),
        base_location=_extract_location(cv_text),
        contact_email=_extract_email(cv_text),
        professional_summary=None,
        roles=_extract_roles_skeleton(cv_text),
        education=[],
        projects=[],
        skills=_extract_skills_inline(cv_text),
        # Phone + LinkedIn + GitHub aren't on CVImport's top level but
        # the LLM tier picks them up. They're still surfaced to the
        # wizard via raw_text searching.
        extraction_confidence=2,
        raw_text=cv_text,
    )


# ---------------------------------------------------------------------------
# Tier 1 — Haiku via call_structured. ~5s.
# ---------------------------------------------------------------------------


async def parse(
    *,
    cv_text: str,
    session_id: Optional[str] = None,
) -> CVImport:
    """Haiku pass: free-form CV text → structured CVImport.

    `cv_text` is the post-extraction raw text — caller is responsible
    for PDF/DOCX → str. The Tier 1 content shield runs here before
    the agent sees any of it (CVs are user input and may carry
    injection attempts pasted from recruiter emails).
    """
    if not cv_text or not cv_text.strip():
        raise ValueError("cv_text is empty")

    shielded = _shield_tier1(cv_text).cleaned_text

    extracted: CVImportLLMOutput = await call_structured(
        agent_name="cv_parser",
        system_prompt=_SYSTEM_PROMPT,
        user_input=shielded,
        output_schema=CVImportLLMOutput,
        model=settings.haiku_model_id,
        effort="medium",
        session_id=session_id,
    )
    out = CVImport(**extracted.model_dump(), raw_text=cv_text)
    logger.info(
        "cv_parser (tier1/Haiku): name=%r roles=%d education=%d projects=%d "
        "skills=%d narrative_chars=%d confidence=%d",
        out.name, len(out.roles), len(out.education),
        len(out.projects), len(out.skills),
        len(out.narrative or ""), out.extraction_confidence,
    )
    return out


# ---------------------------------------------------------------------------
# File-format helpers — caller hands us bytes, we return text
# ---------------------------------------------------------------------------


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "pypdf is required for PDF CV extraction"
        ) from exc

    reader = PdfReader(BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            logger.warning("pypdf page extract failed: %s", exc)
    return "\n\n".join(p for p in pages if p.strip())


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX CV extraction"
        ) from exc

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(*, data: bytes, filename: str) -> str:
    """Dispatch on file extension. Plain-text upload bypasses both."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf_text(data)
    if name.endswith(".docx"):
        return _extract_docx_text(data)
    return data.decode("utf-8", errors="replace")
